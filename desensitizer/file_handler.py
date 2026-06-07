# -*- coding: utf-8 -*-
"""
法律文书脱敏工具 - 文件处理模块
支持格式：PDF、图片(PNG/JPG/BMP/TIFF)、Word(.docx)、文本(.txt)
内置 RapidOCR 引擎（基于PaddleOCR ONNX模型），无需额外安装外部程序
"""

import os
import sys
import tempfile
import threading
import re
from typing import Tuple, Optional
from pathlib import Path

# 在导入onnxruntime之前设置多线程环境变量
_cpu_count = os.cpu_count() or 4
os.environ.setdefault('OMP_NUM_THREADS', str(min(_cpu_count, 8)))
os.environ.setdefault('ORT_THREADS', str(min(_cpu_count, 8)))

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from PIL import Image, ImageOps, ImageFilter, ImageEnhance
except ImportError:
    Image = None
    ImageOps = None
    ImageFilter = None
    ImageEnhance = None

try:
    from rapidocr_onnxruntime import RapidOCR
except ImportError:
    RapidOCR = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import chardet
except ImportError:
    chardet = None


OCR_MAX_SIDE = int(os.environ.get('OCR_MAX_SIDE', '2000'))
OCR_MIN_SIDE = int(os.environ.get('OCR_MIN_SIDE', '1300'))
PDF_OCR_SCALE = float(os.environ.get('PDF_OCR_SCALE', '2.0'))
OCR_MIN_CONFIDENCE = float(os.environ.get('OCR_MIN_CONFIDENCE', '0.35'))


def get_resource_path(relative_path):
    """获取资源文件路径（兼容PyInstaller打包）"""
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', relative_path)


def _median(values):
    values = sorted(v for v in values if v > 0)
    if not values:
        return 0
    mid = len(values) // 2
    if len(values) % 2:
        return values[mid]
    return (values[mid - 1] + values[mid]) / 2


def _box_metrics(box):
    xs = [float(point[0]) for point in box]
    ys = [float(point[1]) for point in box]
    x_min, x_max = min(xs), max(xs)
    y_min, y_max = min(ys), max(ys)
    return {
        'x': x_min,
        'y': y_min,
        'cx': (x_min + x_max) / 2,
        'cy': (y_min + y_max) / 2,
        'width': max(x_max - x_min, 1),
        'height': max(y_max - y_min, 1),
    }


def _looks_cjk_or_number(value):
    return bool(re.search(r'[\u4e00-\u9fff0-9]', value or ''))


def _join_ocr_fragments(fragments):
    """按水平位置合并同一行，避免把身份证号/案号/地址用空格拆坏。"""
    if not fragments:
        return ''

    fragments.sort(key=lambda item: item['x'])
    merged = fragments[0]['text']
    prev = fragments[0]

    for item in fragments[1:]:
        gap = item['x'] - (prev['x'] + prev['width'])
        avg_char_width = max(prev['width'] / max(len(prev['text']), 1), 8)
        prev_tail = merged[-1:] if merged else ''
        next_head = item['text'][:1]

        needs_space = gap > max(avg_char_width * 1.25, 10)
        if _looks_cjk_or_number(prev_tail) or _looks_cjk_or_number(next_head):
            needs_space = False
        if prev_tail in '（([《“' or next_head in '，。；：、,.!?)]）”》':
            needs_space = False

        merged += (' ' if needs_space else '') + item['text']
        prev = item

    return _cleanup_ocr_text(merged)


def _cleanup_ocr_text(text: str) -> str:
    """清理 OCR 常见拼接噪声，提升后续敏感信息规则命中率。"""
    text = re.sub(r'(?<=\d)\s+(?=\d)', '', text)
    text = re.sub(r'(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])', '', text)
    text = re.sub(r'\s+([，。；：、,.!?）)])', r'\1', text)
    text = re.sub(r'([（(])\s+', r'\1', text)
    return text.strip()


class OCREngine:
    """内置OCR引擎 - 基于RapidOCR(ONNX)，无需外部依赖"""

    _instance = None

    @classmethod
    def get_instance(cls):
        """单例模式，避免重复加载模型"""
        if cls._instance is None:
            cls._instance = cls()
        return cls._instance

    def __init__(self):
        self._ocr = None
        self._init_lock = threading.Lock()

    @property
    def is_ready(self):
        return self._ocr is not None

    def _init_ocr(self):
        """延迟初始化OCR引擎。首次识别才加载，避免应用启动时长时间卡住。"""
        if self._ocr is not None:
            return

        with self._init_lock:
            if self._ocr is not None:
                return
            if RapidOCR is None:
                raise ImportError(
                    "RapidOCR 未安装，请运行: pip install rapidocr-onnxruntime"
                )

            self._patch_ort_threading()
            self._ocr = RapidOCR(use_angle_cls=False)

    @staticmethod
    def _patch_ort_threading():
        """修补ONNX Runtime会话选项，启用多线程"""
        try:
            from onnxruntime import SessionOptions, GraphOptimizationLevel, InferenceSession
            from rapidocr_onnxruntime.utils import OrtInferSession

            if getattr(OrtInferSession, '_legal_doc_fast_patch', False):
                return

            cpu_count = min(os.cpu_count() or 4, 8)

            def _fast_init(self, config):
                sess_opt = SessionOptions()
                sess_opt.log_severity_level = 4
                sess_opt.enable_cpu_mem_arena = True
                sess_opt.graph_optimization_level = GraphOptimizationLevel.ORT_ENABLE_ALL
                sess_opt.intra_op_num_threads = cpu_count
                sess_opt.inter_op_num_threads = max(1, cpu_count // 2)

                cpu_ep = 'CPUExecutionProvider'
                ep_list = [(cpu_ep, {'arena_extend_strategy': 'kSameAsRequested'})]

                self._verify_model(config['model_path'])
                self.session = InferenceSession(
                    config['model_path'],
                    sess_options=sess_opt,
                    providers=ep_list
                )

            OrtInferSession.__init__ = _fast_init
            OrtInferSession._legal_doc_fast_patch = True
        except Exception:
            pass

    def _prepare_image(self, image_input):
        """把图片整理成 OCR 友好的尺寸和对比度。"""
        if Image is None:
            raise ImportError("请安装 Pillow: pip install Pillow")

        import numpy as np

        if isinstance(image_input, str):
            img = Image.open(image_input)
            img = ImageOps.exif_transpose(img)
        elif isinstance(image_input, np.ndarray):
            img = Image.fromarray(image_input)
        else:
            return image_input

        if img.mode in ('RGBA', 'LA'):
            background = Image.new('RGB', img.size, 'white')
            alpha = img.getchannel('A') if img.mode == 'RGBA' else img.getchannel(1)
            background.paste(img.convert('RGB'), mask=alpha)
            img = background
        elif img.mode != 'RGB':
            img = img.convert('RGB')

        width, height = img.size
        longest = max(width, height)
        if longest > OCR_MAX_SIDE:
            ratio = OCR_MAX_SIDE / longest
            img = img.resize((max(1, int(width * ratio)), max(1, int(height * ratio))), Image.LANCZOS)
        elif longest < OCR_MIN_SIDE:
            ratio = OCR_MIN_SIDE / longest
            img = img.resize((max(1, int(width * ratio)), max(1, int(height * ratio))), Image.BICUBIC)

        gray = ImageOps.grayscale(img)
        gray = ImageOps.autocontrast(gray)
        gray = ImageEnhance.Contrast(gray).enhance(1.2)
        gray = gray.filter(ImageFilter.SHARPEN)
        return np.array(gray.convert('RGB'))

    def recognize(self, image_input) -> str:
        """
        识别图片中的文字，保留基本换行和缩进。
        """
        self._init_ocr()
        prepared_image = self._prepare_image(image_input)

        result, _ = self._ocr(prepared_image)
        if not result:
            return ''

        fragments = []
        for item in result:
            if len(item) < 3:
                continue
            box, text, confidence = item[0], item[1], item[2]
            try:
                conf_val = float(confidence)
            except (TypeError, ValueError):
                conf_val = 1.0

            text = (text or '').strip()
            if conf_val < OCR_MIN_CONFIDENCE or not text:
                continue

            metrics = _box_metrics(box)
            fragments.append({
                **metrics,
                'text': text,
                'confidence': conf_val,
            })

        if not fragments:
            return ''

        fragments.sort(key=lambda item: (item['cy'], item['x']))
        median_height = _median([item['height'] for item in fragments]) or 18
        y_threshold = max(10, median_height * 0.65)

        rows = []
        current = [fragments[0]]
        current_y = fragments[0]['cy']

        for item in fragments[1:]:
            if abs(item['cy'] - current_y) <= y_threshold:
                current.append(item)
                current_y = sum(part['cy'] for part in current) / len(current)
            else:
                rows.append(current)
                current = [item]
                current_y = item['cy']
        rows.append(current)

        merged_lines = []
        for row in rows:
            avg_y = sum(item['cy'] for item in row) / len(row)
            min_x = min(item['x'] for item in row)
            avg_height = sum(item['height'] for item in row) / len(row)
            text = _join_ocr_fragments(row)
            if text:
                merged_lines.append((avg_y, min_x, text, avg_height))

        if not merged_lines:
            return ''

        base_x = min(line[1] for line in merged_lines)
        output_lines = []
        for i, (y, x, text, height) in enumerate(merged_lines):
            indent = ''
            if x - base_x > 24:
                indent_chars = int((x - base_x) / 18)
                indent = '  ' * min(indent_chars, 6)

            if i > 0:
                prev_y = merged_lines[i - 1][0]
                prev_height = merged_lines[i - 1][3]
                if y - prev_y > prev_height * 1.9:
                    output_lines.append('')

            output_lines.append(indent + text)

        return '\n'.join(output_lines)

    def recognize_from_pil(self, pil_image) -> str:
        """从PIL Image对象识别文字"""
        import numpy as np

        if pil_image.mode != 'RGB':
            pil_image = pil_image.convert('RGB')

        img_array = np.array(pil_image)
        return self.recognize(img_array)


class FileHandler:
    """文件处理器 - 支持多种格式的文本提取"""

    SUPPORTED_EXTENSIONS = {
        'pdf': ['.pdf'],
        'image': ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.tif'],
        'word': ['.docx', '.doc'],
        'text': ['.txt', '.csv', '.log', '.md']
    }

    def __init__(self):
        """初始化文件处理器（OCR引擎内置，无需配置）"""
        self._ocr_engine = None

    @property
    def ocr_engine(self):
        """延迟加载OCR引擎"""
        if self._ocr_engine is None:
            self._ocr_engine = OCREngine.get_instance()
        return self._ocr_engine

    @classmethod
    def get_file_type(cls, filename: str) -> Optional[str]:
        """根据文件扩展名判断文件类型"""
        ext = Path(filename).suffix.lower()
        for file_type, extensions in cls.SUPPORTED_EXTENSIONS.items():
            if ext in extensions:
                return file_type
        return None

    @classmethod
    def is_supported(cls, filename: str) -> bool:
        """检查文件是否支持"""
        return cls.get_file_type(filename) is not None

    def extract_text(self, file_path: str) -> Tuple[str, str]:
        """
        从文件中提取文本
        """
        file_type = self.get_file_type(file_path)
        if file_type is None:
            raise ValueError(f"不支持的文件格式: {Path(file_path).suffix}")

        if file_type == 'pdf':
            return self._extract_from_pdf(file_path), 'pdf'
        if file_type == 'image':
            return self._extract_from_image(file_path), 'image'
        if file_type == 'word':
            return self._extract_from_word(file_path), 'word'
        if file_type == 'text':
            return self._extract_from_text(file_path), 'text'
        raise ValueError(f"不支持的文件类型: {file_type}")

    def extract_text_from_bytes(self, file_bytes: bytes, filename: str) -> Tuple[str, str]:
        """从文件字节流中提取文本"""
        file_type = self.get_file_type(filename)
        if file_type is None:
            raise ValueError(f"不支持的文件格式: {Path(filename).suffix}")

        suffix = Path(filename).suffix
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            text, ftype = self.extract_text(tmp_path)
            return text, ftype
        finally:
            os.unlink(tmp_path)

    def _extract_from_pdf(self, file_path: str) -> str:
        """从PDF文件提取文本。原生文本页走快速路径，扫描页才进入OCR。"""
        if fitz is None:
            raise ImportError("请安装 PyMuPDF: pip install pymupdf")

        text_parts = []
        doc = fitz.open(file_path)

        try:
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text("text") or ''

                if self._page_needs_ocr(page, page_text):
                    ocr_text = self._ocr_pdf_page(page, page_num + 1)
                    if ocr_text.strip():
                        page_text = ocr_text

                if page_text.strip():
                    text_parts.append(f"--- 第 {page_num + 1} 页 ---\n{page_text.strip()}")
        finally:
            doc.close()

        return '\n\n'.join(text_parts)

    @staticmethod
    def _page_needs_ocr(page, page_text: str) -> bool:
        stripped = (page_text or '').strip()
        if len(stripped) >= 20:
            return False
        try:
            return bool(page.get_images(full=True)) or len(stripped) < 10
        except Exception:
            return len(stripped) < 10

    def _ocr_pdf_page(self, page, page_number: int) -> str:
        """以较高倍率渲染扫描页，再交给 OCR。"""
        if fitz is None:
            raise ImportError("请安装 PyMuPDF: pip install pymupdf")

        import numpy as np

        try:
            matrix = fitz.Matrix(PDF_OCR_SCALE, PDF_OCR_SCALE)
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
                pix.height, pix.width, pix.n
            )
            if pix.n == 1:
                img_array = np.repeat(img_array, 3, axis=2)
            elif pix.n > 3:
                img_array = img_array[:, :, :3]
            return self.ocr_engine.recognize(img_array)
        except Exception as exc:
            print(f"[OCR] 第 {page_number} 页识别失败: {exc}")
            return ''

    def _extract_from_image(self, file_path: str) -> str:
        """从图片文件提取文本（内置OCR）"""
        if Image is None:
            raise ImportError("请安装 Pillow: pip install Pillow")

        return self.ocr_engine.recognize(file_path)

    def _extract_from_word(self, file_path: str) -> str:
        """从Word文件提取文本（保留段落格式）"""
        if Document is None:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = Document(file_path)
        text_parts = []

        for para in doc.paragraphs:
            text_parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return '\n'.join(text_parts)

    def _extract_from_text(self, file_path: str) -> str:
        """从纯文本文件提取文本"""
        with open(file_path, 'rb') as f:
            raw_data = f.read()

        if chardet:
            detected = chardet.detect(raw_data)
            encoding = detected.get('encoding', 'utf-8')
        else:
            encoding = 'utf-8'

        try:
            return raw_data.decode(encoding)
        except (UnicodeDecodeError, TypeError):
            return raw_data.decode('utf-8', errors='replace')


class PDFRedactor:
    """PDF脱敏输出器 - 生成脱敏后的PDF"""

    @staticmethod
    def create_redacted_pdf(original_path: str, redacted_text: str, output_path: str):
        """创建脱敏后的PDF文件"""
        if fitz is None:
            raise ImportError("请安装 PyMuPDF: pip install pymupdf")

        doc = fitz.open()
        pages = redacted_text.split('--- 第')

        for page_content in pages:
            if not page_content.strip():
                continue

            lines = page_content.split('\n')
            if lines and '页 ---' in lines[0]:
                lines = lines[1:]

            content = '\n'.join(lines).strip()
            if not content:
                continue

            page = doc.new_page(width=595, height=842)
            text_writer = fitz.TextWriter(page.rect)
            font = fitz.Font("china-s")

            y_pos = 50
            for line in content.split('\n'):
                if y_pos > 790:
                    page = doc.new_page(width=595, height=842)
                    text_writer = fitz.TextWriter(page.rect)
                    y_pos = 50

                try:
                    text_writer.append((50, y_pos), line, font=font, fontsize=11)
                except Exception:
                    text_writer.append((50, y_pos), line.encode('ascii', 'replace').decode(), fontsize=11)
                y_pos += 18

            text_writer.write_text(page)

        doc.save(output_path)
        doc.close()

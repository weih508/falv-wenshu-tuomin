# -*- coding: utf-8 -*-
"""
法律文书脱敏工具 - Flask Web应用主程序
支持PyInstaller打包为独立exe
"""

import os
import sys
import uuid
import time
import webbrowser
import threading
from pathlib import Path

# 无控制台模式下（console=False），stdout/stderr为None会导致print崩溃
# 需要重定向到日志文件或devnull
if getattr(sys, 'frozen', False) and (sys.stdout is None or sys.stderr is None):
    _log_dir = Path(os.path.dirname(sys.executable))
    _log_file = open(str(_log_dir / 'app.log'), 'w', encoding='utf-8')
    if sys.stdout is None:
        sys.stdout = _log_file
    if sys.stderr is None:
        sys.stderr = _log_file

from flask import Flask, render_template, request, jsonify, send_file


def get_base_path():
    """获取基础路径（兼容打包环境）"""
    if getattr(sys, 'frozen', False):
        return Path(sys._MEIPASS)
    return Path(__file__).parent


def get_work_dir():
    """获取工作目录（存放上传和输出文件）"""
    if getattr(sys, 'frozen', False):
        return Path(os.path.dirname(sys.executable))
    return Path(__file__).parent


BASE_PATH = get_base_path()
WORK_DIR = get_work_dir()
UPLOAD_FOLDER = WORK_DIR / 'uploads'
OUTPUT_FOLDER = WORK_DIR / 'outputs'
UPLOAD_FOLDER.mkdir(exist_ok=True)
OUTPUT_FOLDER.mkdir(exist_ok=True)

# 添加模块搜索路径
if getattr(sys, 'frozen', False):
    sys.path.insert(0, str(BASE_PATH))

from desensitizer import SensitiveDetector, FileHandler, PDFRedactor

ALLOWED_EXTENSIONS = {
    '.pdf', '.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.tif',
    '.docx', '.doc', '.txt', '.csv', '.md'
}

app = Flask(__name__,
            template_folder=str(BASE_PATH / 'templates'),
            static_folder=str(BASE_PATH / 'static'))
app.config['UPLOAD_FOLDER'] = str(UPLOAD_FOLDER)
app.config['OUTPUT_FOLDER'] = str(OUTPUT_FOLDER)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

# 初始化核心组件
detector = SensitiveDetector()
file_handler = FileHandler()


def allowed_file(filename):
    """检查文件是否允许上传"""
    return Path(filename).suffix.lower() in ALLOWED_EXTENSIONS


def warmup_ocr_async():
    """后台预热 OCR，避免程序启动时长时间卡住。"""
    try:
        print("  [OCR] 后台加载识别模型中，界面可先使用文本/Word/PDF原生文字...")
        file_handler.ocr_engine._init_ocr()
        print("  [OCR] 模型加载完成。")
    except Exception as exc:
        print(f"  [OCR] 模型加载失败: {exc}（图片/扫描PDF识别功能可能不可用）")


@app.route('/')
def index():
    """主页"""
    return render_template('index.html')


@app.route('/api/categories', methods=['GET'])
def get_categories():
    """获取所有脱敏类别"""
    categories = detector.get_categories()
    return jsonify({'success': True, 'categories': categories})


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文件并提取文本"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': '未选择文件'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'error': '文件名为空'}), 400

    if not allowed_file(file.filename):
        return jsonify({
            'success': False,
            'error': f'不支持的文件格式，支持: {", ".join(sorted(ALLOWED_EXTENSIONS))}'
        }), 400

    original_ext = Path(file.filename).suffix.lower()
    unique_name = f"{uuid.uuid4().hex}{original_ext}"
    file_path = UPLOAD_FOLDER / unique_name
    file.save(str(file_path))

    try:
        text, file_type = file_handler.extract_text(str(file_path))

        if not text.strip():
            return jsonify({
                'success': False,
                'error': '无法从文件中提取文本。文件可能为空、格式损坏，或OCR引擎不可用。'
            }), 400

        return jsonify({
            'success': True,
            'text': text,
            'file_type': file_type,
            'file_id': unique_name,
            'original_name': file.filename
        })

    except ImportError as e:
        return jsonify({'success': False, 'error': str(e)}), 500
    except Exception as e:
        return jsonify({'success': False, 'error': f'文件处理失败: {str(e)}'}), 500


@app.route('/api/detect', methods=['POST'])
def detect_sensitive():
    """检测敏感信息"""
    data = request.get_json()
    if not data or 'text' not in data:
        return jsonify({'success': False, 'error': '缺少文本内容'}), 400

    text = data['text']
    categories = data.get('categories', None)
    results = detector.detect(text, categories)

    return jsonify({
        'success': True,
        'results': results,
        'total': len(results)
    })


@app.route('/api/desensitize', methods=['POST'])
def desensitize():
    """执行脱敏处理"""
    data = request.get_json()
    if not data or 'text' not in data:
        return jsonify({'success': False, 'error': '缺少文本内容'}), 400

    text = data['text']
    categories = data.get('categories', None)
    selected_items = data.get('selected_items', None)

    desensitized_text, results = detector.desensitize(text, categories, selected_items)

    return jsonify({
        'success': True,
        'original_text': text,
        'desensitized_text': desensitized_text,
        'results': results,
        'total_detected': len(results),
        'total_masked': len(selected_items) if selected_items else len(results)
    })


@app.route('/api/export', methods=['POST'])
def export_file():
    """导出脱敏后的文件"""
    data = request.get_json()
    if not data or 'text' not in data:
        return jsonify({'success': False, 'error': '缺少文本内容'}), 400

    text = data['text']
    export_format = data.get('format', 'txt')
    original_name = data.get('original_name', 'desensitized')

    base_name = Path(original_name).stem
    timestamp = int(time.time())

    if export_format == 'txt':
        output_name = f"{base_name}_脱敏_{timestamp}.txt"
        output_path = OUTPUT_FOLDER / output_name
        with open(str(output_path), 'w', encoding='utf-8-sig') as f:
            f.write(text)
    elif export_format == 'docx':
        try:
            from docx import Document as DocxDocument
            output_name = f"{base_name}_脱敏_{timestamp}.docx"
            output_path = OUTPUT_FOLDER / output_name

            file_id = data.get('file_id', '')
            original_docx_path = UPLOAD_FOLDER / file_id if file_id else None

            if original_docx_path and original_docx_path.exists() and file_id.endswith('.docx'):
                doc = DocxDocument(str(original_docx_path))
                for para in doc.paragraphs:
                    for run in para.runs:
                        new_text = run.text
                        for item in data.get('results', []):
                            if item.get('text', '') in new_text:
                                new_text = new_text.replace(item['text'], item['masked'])
                        run.text = new_text
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    new_text = run.text
                                    for item in data.get('results', []):
                                        if item.get('text', '') in new_text:
                                            new_text = new_text.replace(item['text'], item['masked'])
                                    run.text = new_text
                doc.save(str(output_path))
            else:
                doc = DocxDocument()
                for para in text.split('\n'):
                    doc.add_paragraph(para)
                doc.save(str(output_path))
        except ImportError:
            return jsonify({'success': False, 'error': '需要安装 python-docx'}), 500
    elif export_format == 'md':
        output_name = f"{base_name}_脱敏_{timestamp}.md"
        output_path = OUTPUT_FOLDER / output_name
        with open(str(output_path), 'w', encoding='utf-8') as f:
            f.write(text)
    elif export_format == 'pdf':
        try:
            output_name = f"{base_name}_脱敏_{timestamp}.pdf"
            output_path = OUTPUT_FOLDER / output_name
            PDFRedactor.create_redacted_pdf(None, text, str(output_path))
        except Exception as e:
            return jsonify({'success': False, 'error': f'PDF导出失败: {str(e)}'}), 500
    else:
        return jsonify({'success': False, 'error': f'不支持的导出格式: {export_format}'}), 400

    return jsonify({
        'success': True,
        'filename': output_name,
        'download_url': f'/api/download/{output_name}'
    })


@app.route('/api/download/<filename>', methods=['GET'])
def download_file(filename):
    """下载文件"""
    safe_filename = Path(filename).name
    file_path = OUTPUT_FOLDER / safe_filename
    if not file_path.exists():
        return jsonify({'success': False, 'error': '文件不存在'}), 404
    return send_file(str(file_path), as_attachment=True)


@app.route('/api/paste-text', methods=['POST'])
def paste_text():
    """直接粘贴文本进行检测"""
    data = request.get_json()
    if not data or 'text' not in data:
        return jsonify({'success': False, 'error': '缺少文本内容'}), 400

    return jsonify({
        'success': True,
        'text': data['text'],
        'file_type': 'text',
        'file_id': None,
        'original_name': '粘贴文本'
    })


@app.route('/api/ocr-status', methods=['GET'])
def ocr_status():
    """检查OCR引擎状态"""
    try:
        from rapidocr_onnxruntime import RapidOCR  # noqa: F401
        engine = file_handler.ocr_engine
        return jsonify({
            'success': True,
            'engine': 'RapidOCR',
            'ready': engine.is_ready,
            'status': '已加载' if engine.is_ready else '可用，首次识别时自动加载'
        })
    except ImportError:
        return jsonify({'success': False, 'engine': 'RapidOCR', 'status': '未安装'})


# 清理过期文件
def cleanup_old_files(max_age_hours=24):
    """清理超过指定时间的上传文件"""
    current_time = time.time()
    for folder in [UPLOAD_FOLDER, OUTPUT_FOLDER]:
        for file_path in folder.iterdir():
            if file_path.is_file():
                file_age = current_time - file_path.stat().st_mtime
                if file_age > max_age_hours * 3600:
                    file_path.unlink()


if __name__ == '__main__':
    port = 5000
    print("=" * 60)
    print("  法律文书脱敏审查工具")
    print(f"  启动地址: http://127.0.0.1:{port}")
    print("  按 Ctrl+C 停止服务")
    print("=" * 60)

    cleanup_old_files()
    threading.Thread(target=warmup_ocr_async, daemon=True).start()

    print("  [启动完成] 正在打开浏览器...")
    print("=" * 60)

    def open_browser():
        webbrowser.open(f'http://127.0.0.1:{port}')

    threading.Timer(0.5, open_browser).start()

    app.run(host='127.0.0.1', port=port, debug=False)
